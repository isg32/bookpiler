import os
import re
import logging
import fitz  # PyMuPDF
import unicodedata
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Setup logger
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
logger = logging.getLogger("Bookpiler")

# Config
DATA_DIR = './data'
ASSET_DIR = './assets'
HEADER_LOGO = os.path.join(ASSET_DIR, 'amarujalalogo.png')
ALLOWED_EXT = ['.txt', '.pdf']

# --- Helper Functions ---
def clean_line(line: str) -> str:
    return ''.join(c for c in line if unicodedata.category(c)[0] != 'C').strip()

def read_file_text(path):
    if path.endswith('.txt'):
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    elif path.endswith('.pdf'):
        doc = fitz.open(path)
        return "\n".join([page.get_text() for page in doc])
    return ""

def extract_chapter_title_from_text(text):
    lines = text.strip().splitlines()
    for line in lines:
        clean = clean_line(line)
        if clean.lower().startswith("chapter"):
            return clean
    return "Untitled Chapter"

def extract_chapter_number(title):
    match = re.search(r'chapter\s*(\d+)', title.lower())
    return int(match.group(1)) if match else 9999

def render_text_block(doc, text):
    lines = text.strip().splitlines()
    for line in lines:
        line = clean_line(line)
        if not line:
            continue
        if line.isdigit() and len(line) < 4:
            logger.debug(f"Skipping probable page number: {line}")
            continue
        if line.startswith('[image:') and line.endswith(']'):
            img_path = line[7:-1].strip()
            if os.path.exists(img_path):
                logger.info(f"Adding image: {img_path}")
                doc.add_picture(img_path, width=Inches(4.5))
            else:
                logger.warning(f"Image not found: {img_path}")
        else:
            doc.add_paragraph(line)

def add_separator(doc):
    para = doc.add_paragraph()
    run = para.add_run("_" * 70)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def insert_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_header_footer(section, header_text):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    header_para.clear()
    run = header_para.add_run()

    if os.path.exists(HEADER_LOGO):
        try:
            run.add_picture(HEADER_LOGO, width=Inches(0.8))
        except Exception as e:
            logger.warning(f"Logo insert failed: {e}")

    header_para.add_run(f"   {header_text}").bold = True
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    footer_para.clear()
    insert_page_number(footer_para)
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# --- Folder Parser ---
def parse_structure(folders):
    logger.info("📁 Parsing selected folders...")
    book = {}
    for class_folder in folders:
        class_path = os.path.join(DATA_DIR, class_folder)
        if not os.path.isdir(class_path):
            continue

        files = [f for f in os.listdir(class_path) if os.path.splitext(f)[1].lower() in ALLOWED_EXT]
        for file in files:
            lower_file = file.lower()
            full_path = os.path.join(class_path, file)
            try:
                parts = file.split(' - ')
                chapter_key = parts[2].strip()
            except:
                logger.warning(f"Skipping malformed filename: {file}")
                continue

            if 'explanation' in lower_file:
                book.setdefault(class_folder, {}).setdefault(chapter_key, {})['explanation'] = full_path
            elif 'question' in lower_file:
                book.setdefault(class_folder, {}).setdefault(chapter_key, {})['questions'] = full_path

    return book

# --- Book Generator ---
def create_book_for_folder(folder_name, chapters_dict):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    first_section = True

    class_info = folder_name.split()
    class_name = class_info[1]
    subject = class_info[2]

    sorted_chapters = sorted(
        chapters_dict.items(),
        key=lambda item: extract_chapter_number(
            extract_chapter_title_from_text(read_file_text(item[1].get('questions', '')))
        )
    )

    for chapter_key, chapter_files in sorted_chapters:
        explanation_text = read_file_text(chapter_files.get('explanation', ''))
        questions_text = read_file_text(chapter_files.get('questions', ''))

        chapter_title = extract_chapter_title_from_text(questions_text or explanation_text or chapter_key)
        logger.info(f"📝 Writing: {chapter_title}")

        section = doc.sections[0] if first_section else doc.add_section()
        first_section = False

        header_text = f"Class {class_name}, {subject} - {chapter_title}"
        add_header_footer(section, header_text)

        doc.add_page_break()

        # Chapter Title
        para = doc.add_paragraph()
        run = para.add_run(chapter_title)
        run.bold = True
        run.font.size = Pt(24)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        add_separator(doc)

        if explanation_text:
            doc.add_heading("Explanations", level=2)
            lines = explanation_text.strip().splitlines()
            if lines and "explanation" in clean_line(lines[0]).lower():
                lines = lines[1:]
            render_text_block(doc, "\n".join(lines))

        if questions_text:
            doc.add_heading("Questions", level=2)
            lines = questions_text.strip().splitlines()
            if lines and clean_line(lines[0]).lower().startswith("chapter"):
                lines = lines[1:]
            render_text_block(doc, "\n".join(lines))

    os.makedirs("generated", exist_ok=True)
    output_path = f"./generated/{folder_name} - Compiled Book.docx"
    logger.info(f"💾 Saving to: {output_path}")
    doc.save(output_path)
    logger.info("✅ Book compiled successfully for: " + folder_name)

# --- Entry Point ---
if __name__ == "__main__":
    logger.info("📘 Bookpiler started")

    all_folders = [f for f in os.listdir(DATA_DIR) if os.path.isdir(os.path.join(DATA_DIR, f))]
    if not all_folders:
        logger.error("❌ No class folders found in ./data/")
        exit(1)

    print("\nFolders found. Choose to process:")
    print("[ 0 ] All")
    for idx, name in enumerate(all_folders, 1):
        print(f"[ {idx} ] {name}")

    choice = input("\nEnter your choice (number): ").strip()

    if not choice.isdigit() or int(choice) < 0 or int(choice) > len(all_folders):
        logger.error("❌ Invalid choice. Exiting.")
        exit(1)

    choice = int(choice)
    selected_folders = all_folders if choice == 0 else [all_folders[choice - 1]]

    for folder in selected_folders:
        logger.info(f"📚 Processing: {folder}")
        book_data = parse_structure([folder])
        if folder in book_data:
            create_book_for_folder(folder, book_data[folder])
        else:
            logger.warning(f"No valid chapters found in {folder}.")

    logger.info("📗 Done.")
